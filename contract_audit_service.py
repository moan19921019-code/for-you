# -*- coding: utf-8 -*-
import os, re, datetime, tempfile, subprocess, json
from flask import Flask, request, jsonify, render_template_string
from flask_cors import CORS
from io import BytesIO

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024
CORS(app)

# ===== 审核配置（从 JSON 文件加载） =====
CONFIG_FILE = os.path.join(os.path.dirname(__file__), 'audit_rules.json')
DEFAULT_CONFIG = {
    "基本信息": {
        "乙方名称": "",
        "乙方地址": "",
        "乙方电话": "",
        "收款户名": "",
        "收款开户行": "",
        "收款账号": ""
    },
    "开户行变体": [],
    "条款检查": {
        "必须有违约责任条款": True,
        "必须有保密条款": True,
        "必须有知识产权条款": True,
        "必须有争议解决条款": True,
        "必须有付款条款": True
    }
}

def load_config():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            pass
    return DEFAULT_CONFIG.copy()

def save_config(data):
    with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

# 首次启动自动创建配置文件
if not os.path.exists(CONFIG_FILE):
    save_config(DEFAULT_CONFIG)

# 加载配置到全局变量
CFG = load_config()
STANDARDS = CFG.get("基本信息", DEFAULT_CONFIG["基本信息"])
BANK_VARIANTS = CFG.get("开户行变体", DEFAULT_CONFIG["开户行变体"])

def extract_text(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.docx':
        return extract_docx(filepath)
    elif ext == '.pdf':
        return extract_pdf(filepath)
    elif ext == '.doc':
        return extract_doc(filepath)
    return "[ERROR] 不支持格式: " + ext

def extract_docx(filepath):
    try:
        import docx
        doc = docx.Document(filepath)
        return '\n'.join(p.text for p in doc.paragraphs)
    except Exception as e:
        return "[ERROR docx] " + str(e)

def _clean_ocr_text(text):
    """OCR后处理：去空格、去噪音字符、修复常见OCR错误"""
    import re
    # 1. 去掉连续空格
    text = re.sub(r' {2,}', ' ', text)
    # 2. 去掉相邻非空字符之间的空格（Tesseract会将每个字符分开识别）
    text = re.sub(r'(?<=\S) (?=\S)', '', text)
    # 3. 去掉行尾和文件末尾的OCR噪音字符
    text = re.sub(r'[_\|\x60\u201c\u201d\u300e\u300f\u3010\u3011\u2018\u2019「」『』]+(?=\n|$)', '', text)
    text = re.sub(r'^[_\|\x60\u201c\u201d]+', '', text, flags=re.MULTILINE)
    # 4. 去掉孤立噪音字符行（纯噪音字符组成的短行，如表格边框识别产物）
    text = re.sub(r'^[a-zA-Z\u4e00-\u9fff]{1,3}[^a-zA-Z\u4e00-\u9fff\n]{5,}[a-zA-Z\u4e00-\u9fff]{1,3}$', '', text, flags=re.MULTILINE)
    # 5. 修复常见OCR误识别
    fixes = [
        # 司/口 混淆
        (r'信息科技有限公口(?!\))', '信息科技有限公司'),
        # 电话/金额后面的下划线
        (r'(\d)[_\|]?(?=\n|$)', r'\1'),
        # 句号前多余字符
        (r'[a-zA-Z\x60]+(?=[，。；：、])', ''),
        # 数字0/字母O混淆（针对账号）
        (r'([\d]{15,20})', lambda m: m.group(1)),
    ]
    for pat, repl in fixes:
        text = re.sub(pat, repl, text)
    # 6. 去掉纯标点/乱码的孤立行
    lines = text.split('\n')
    cleaned = []
    for line in lines:
        line = line.strip()
        if not line:
            cleaned.append('')
            continue
        # 保留有实质内容的行（中文、英文、数字混合至少有一种超过2个）
        has_cjk = sum(1 for c in line if '\u4e00' <= c <= '\u9fff') >= 2
        has_latin = sum(1 for c in line if c.isalpha()) >= 2
        has_digit = sum(1 for c in line if c.isdigit()) >= 2
        if has_cjk or has_latin or has_digit:
            cleaned.append(line)
        else:
            cleaned.append('')  # 噪音行置空
    return '\n'.join(cleaned)


def extract_pdf(filepath):
    try:
        import fitz
        doc = fitz.open(filepath)
        text = ''.join(page.get_text() for page in doc)
        if text.strip():
            return text

        # 扫描件：用 PyMuPDF 渲染页面 + Tesseract OCR
        tesseract_path = None
        for p in [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            'tesseract',
        ]:
            if os.path.exists(p):
                tesseract_path = p
                break
            if p == 'tesseract':
                try:
                    r = subprocess.run(['tesseract', '--version'], capture_output=True, timeout=5)
                    if r.returncode == 0:
                        tesseract_path = 'tesseract'
                        break
                except:
                    pass

        if not tesseract_path:
            return "[WARNING] PDF 为扫描件，且未检测到 Tesseract OCR，无法识别"

        # 渲染每一页为 PNG，用 Tesseract OCR 识别
        ocr_texts = []
        for i, page in enumerate(doc):
            mat = fitz.Matrix(2.5, 2.5)  # 2.5倍分辨率，提升识别精度
            clip = page.rect
            pix = page.get_pixmap(matrix=mat, clip=clip)
            img_bytes = pix.tobytes("png")

            # 写入临时文件供 tesseract 读取
            fd, img_path = tempfile.mkstemp(suffix='.png')
            os.close(fd)
            try:
                with open(img_path, 'wb') as f:
                    f.write(img_bytes)

                proc = subprocess.run([
                    tesseract_path, img_path, 'stdout',
                    '-l', 'chi_sim+eng',
                    '--psm', '6',
                    '--oem', '3',
                ], capture_output=True, timeout=60)
                page_text = proc.stdout.decode('utf-8', errors='ignore').strip()
                if page_text:
                    page_text = _clean_ocr_text(page_text)
                    ocr_texts.append(page_text)
            finally:
                try: os.remove(img_path)
                except: pass

        if ocr_texts:
            return '[OCR] 扫描件，通过 Tesseract OCR 识别\n\n' + '\n\n'.join(ocr_texts)
        return "[WARNING] PDF 为扫描件，OCR 识别结果为空"

    except Exception as e:
        return "[ERROR pdf] " + str(e)

def extract_doc(filepath):
    """Extract text from .doc file. Tries multiple methods in order:
    1. python-docx (some .doc files are actually .docx format)
    2. LibreOffice soffice (best quality for Chinese)
    3. Windows Word COM (requires Microsoft Word)
    4. olefile + GBK decoding (pure Python, handles Chinese)
    5. Improved binary reader with encoding detection
    """

    # Method 1: Try python-docx first (some .doc files are secretly .docx)
    try:
        import docx
        docx.Document(filepath)
        return extract_docx(filepath)
    except Exception:
        pass

    # Method 2: LibreOffice
    text = soffice_convert(filepath)
    if text and not text.startswith("[ERROR"):
        return text

    # Method 3: Windows Word COM via PowerShell
    text = extract_doc_via_word_com(filepath)
    if text and not text.startswith("[ERROR"):
        return text

    # Method 4: olefile-based extraction (handles OLE2 + Chinese encoding)
    text = extract_doc_via_olefile(filepath)
    if text and not text.startswith("[ERROR"):
        return text

    # Method 5: Improved binary reader with GBK/UTF-16LE detection
    text = read_doc_binary_v2(filepath)
    if text and not text.startswith("[ERROR"):
        return text

    return "[ERROR] 无法解析 .doc 文件。请用 Word 另存为 .docx 格式后重新上传。"


def extract_doc_via_word_com(filepath):
    """Use Windows Word COM object to extract text. Requires Microsoft Word."""
    import subprocess
    abs_path = os.path.abspath(filepath).replace('\\', '\\\\')
    ps_script = f'''
$ErrorActionPreference = 'Stop'
try {{
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $word.DisplayAlerts = 0
    $doc = $word.Documents.Open('{abs_path}', $false, $true, $false)
    $text = $doc.Content.Text
    $doc.Close($false)
    $word.Quit()
    [Console]::OutputEncoding = [Text.Encoding]::UTF8
    Write-Output $text
}} catch {{
    Write-Output "[ERROR] Word COM 提取失败: $_"
}}
'''
    try:
        proc = subprocess.run(
            ['powershell', '-NoProfile', '-ExecutionPolicy', 'Bypass', '-Command', ps_script],
            capture_output=True, timeout=30,
            encoding='utf-8', errors='replace'
        )
        text = proc.stdout.strip()
        if text and not text.startswith('[ERROR]') and len(text) > 20:
            return text
        return "[ERROR] Word COM 未返回有效内容"
    except subprocess.TimeoutExpired:
        return "[ERROR] Word COM 提取超时"
    except Exception as e:
        return f"[ERROR] Word COM: {e}"


def extract_doc_via_olefile(filepath):
    """Extract text from OLE2 .doc using olefile package.
    Handles Chinese text stored as UTF-16LE or GBK in the WordDocument stream.
    """
    try:
        import olefile
    except ImportError:
        return "[ERROR] olefile 未安装 (pip install olefile)"

    try:
        ole = olefile.OleFileIO(filepath)
        if not ole.exists('WordDocument'):
            ole.close()
            return "[ERROR] 不是有效的 Word OLE2 文档"

        # Read the WordDocument stream
        doc_stream = ole.openstream('WordDocument').read()

        # The FIB (File Information Block) starts at offset 0
        # flags at offset 0x000A tells us if text is Unicode or ASCII
        if len(doc_stream) < 0x000C:
            ole.close()
            return "[ERROR] WordDocument 流太短"

        # Check the fComplex flag (bit 3 of flags at offset 0x000A)
        flags = doc_stream[0x000A]
        is_complex = bool(flags & 0x04)

        # ccpText at offset 0x004C (4 bytes) — number of chars in main text
        if len(doc_stream) < 0x0050:
            ole.close()
            return "[ERROR] 无法读取文本长度"

        import struct
        ccp_text = struct.unpack_from('<I', doc_stream, 0x004C)[0]

        if ccp_text == 0 or ccp_text > 1000000:
            ole.close()
            return "[ERROR] 文本长度为0或异常"

        # Try to read from 1Table or 0Table stream which contains the actual text
        table_name = '1Table' if ole.exists('1Table') else '0Table'
        table_stream = ole.openstream(table_name).read()

        # Extract text: try UTF-16LE first (Unicode .doc), then GBK (old Chinese .doc)
        results = []
        extracted = ''

        # Try UTF-16LE decoding — most Chinese .doc files use this
        # The text is stored as UTF-16LE in the table stream
        # We scan for runs of valid UTF-16LE characters
        utf16_text = extract_utf16le_text(table_stream)
        if utf16_text and len(utf16_text) > 20:
            extracted = utf16_text
        else:
            # Try GBK/GB2312 decoding
            extracted = extract_gbk_text(table_stream)

        ole.close()

        if extracted and len(extracted) > 20:
            return extracted
        return "[ERROR] olefile 未能提取到有效文本"

    except Exception as e:
        return f"[ERROR] olefile: {e}"


def extract_utf16le_text(data):
    """Extract Chinese text from binary data by detecting UTF-16LE sequences."""
    result = []
    i = 0
    consecutive_cjk = 0
    while i < len(data) - 1:
        # Read a UTF-16LE code unit
        code = data[i] | (data[i+1] << 8)
        char = None
        if 0x4E00 <= code <= 0x9FFF or 0x3400 <= code <= 0x4DBF:  # CJK Unified
            char = chr(code)
            consecutive_cjk += 1
        elif 0x3000 <= code <= 0x303F:  # CJK punctuation
            char = chr(code)
        elif 0xFF00 <= code <= 0xFFEF:  # Fullwidth forms
            char = chr(code)
        elif 0x0020 <= code <= 0x007E:  # ASCII
            char = chr(code)
            consecutive_cjk = 0
        elif code == 0x000D or code == 0x000A:  # CR/LF
            char = chr(code)
            consecutive_cjk = 0
        else:
            # Only break on truly invalid sequences
            if consecutive_cjk < 2 and not result:
                char = None
            else:
                char = None

        if char is not None:
            result.append(char)
            i += 2
        else:
            i += 1

    text = ''.join(result)
    # Clean up: collapse multiple spaces and newlines
    text = re.sub(r' {2,}', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def extract_gbk_text(data):
    """Try to extract Chinese text by decoding as GBK/GB2312."""
    result = bytearray()
    for byte in data:
        if 32 <= byte <= 126 or byte in (9, 10, 13):
            result.append(byte)
        elif 0x81 <= byte <= 0xFE:  # GBK first byte range
            result.append(byte)

    # Try different encodings
    for enc in ['gb18030', 'gbk', 'gb2312']:
        try:
            text = bytes(result).decode(enc, errors='replace')
            # Count CJK characters to validate
            cjk_count = sum(1 for c in text if '一' <= c <= '鿿')
            if cjk_count > 10:
                # Clean up replacement characters
                text = text.replace('�', '')
                text = re.sub(r' {2,}', ' ', text)
                text = re.sub(r'\n{3,}', '\n\n', text)
                return text.strip()
        except Exception:
            continue
    return ''


def soffice_convert(filepath):
    """Convert .doc to text using LibreOffice headless mode."""
    soffice_paths = [
        r'C:\Program Files\LibreOffice\program\soffice.exe',
        r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
        'soffice',
    ]
    soffice = None
    for p in soffice_paths:
        if os.path.exists(p):
            soffice = p
            break
        if p == 'soffice':
            try:
                r = subprocess.run(['soffice', '--version'], capture_output=True, timeout=5)
                if r.returncode == 0:
                    soffice = 'soffice'
                    break
            except:
                pass
    if not soffice:
        return "[ERROR] LibreOffice 未安装，无法处理 .doc 文件"

    tmpdir = tempfile.gettempdir()
    try:
        subprocess.run([
            soffice, '--headless', '--convert-to', 'txt',
            '--outdir', tmpdir, filepath
        ], capture_output=True, timeout=30)
        base = os.path.splitext(os.path.basename(filepath))[0]
        for fname in os.listdir(tmpdir):
            if fname == base + '.txt':
                path = os.path.join(tmpdir, fname)
                for enc in ['gbk', 'gb18030', 'utf-8']:
                    try:
                        with open(path, 'r', encoding=enc) as f:
                            content = f.read()
                        if content and sum(1 for c in content if '一' <= c <= '鿿') > 10:
                            os.remove(path)
                            return content
                    except:
                        pass
                os.remove(path)
                return "[ERROR] LibreOffice 转换成功但无法读取内容"
    except subprocess.TimeoutExpired:
        return "[ERROR] LibreOffice 转换超时"
    except Exception as e:
        return "[ERROR] " + str(e)
    return "[ERROR] LibreOffice 转换未生成文件"


def read_doc_binary_v2(filepath):
    """Improved binary reader for .doc files with encoding detection."""
    try:
        with open(filepath, 'rb') as f:
            data = f.read()

        # Check OLE2 magic
        if data[:8] != b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':
            return "[ERROR] 不是有效的 Word 文档"

        # Strategy: extract all byte sequences that look like text in common Chinese encodings
        # Try UTF-16LE first
        utf16_result = extract_utf16le_text(data)
        if utf16_result and len(utf16_result) > 50:
            return utf16_result

        # Try GBK decoding
        gbk_result = extract_gbk_text(data)
        if gbk_result and len(gbk_result) > 50:
            return gbk_result

        # Last resort: basic filtering that preserves multi-byte sequences
        result = bytearray()
        for byte in data:
            if 32 <= byte <= 126 or byte in (9, 10, 13):
                result.append(byte)
            elif byte >= 0x80:  # Keep high bytes (might be part of Chinese chars)
                result.append(byte)

        # Try to decode as GB18030
        try:
            text = bytes(result).decode('gb18030', errors='replace')
            text = text.replace('�', '')
            cjk_count = sum(1 for c in text if '一' <= c <= '鿿')
            if cjk_count > 10:
                text = re.sub(r' {2,}', ' ', text)
                text = re.sub(r'\n{3,}', '\n\n', text)
                return text.strip()
        except Exception:
            pass

        return "[ERROR] 无法解码文档内容，请用 Word 另存为 .docx 后重新上传"

    except Exception as e:
        return f"[ERROR] {e}"
    soffice_paths = [
        r'C:\Program Files\LibreOffice\program\soffice.exe',
        r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
        'soffice',
    ]
    soffice = None
    for p in soffice_paths:
        if os.path.exists(p):
            soffice = p
            break
        if p == 'soffice':
            try:
                r = subprocess.run(['soffice', '--version'], capture_output=True, timeout=5)
                if r.returncode == 0:
                    soffice = 'soffice'
                    break
            except:
                pass
    if not soffice:
        return "[ERROR] LibreOffice 未安装，无法处理 .doc 文件"

    tmpdir = tempfile.gettempdir()
    try:
        subprocess.run([
            soffice, '--headless', '--convert-to', 'txt',
            '--outdir', tmpdir, filepath
        ], capture_output=True, timeout=30)
        base = os.path.splitext(os.path.basename(filepath))[0]
        for fname in os.listdir(tmpdir):
            if fname == base + '.txt':
                path = os.path.join(tmpdir, fname)
                for enc in ['gbk', 'gb18030', 'utf-8']:
                    try:
                        with open(path, 'r', encoding=enc) as f:
                            content = f.read()
                        if content and sum(1 for c in content if '\u4e00' <= c <= '\u9fff') > 10:
                            os.remove(path)
                            return content
                    except:
                        pass
                os.remove(path)
                return "[ERROR] LibreOffice 转换成功但无法读取内容"
    except subprocess.TimeoutExpired:
        return "[ERROR] LibreOffice 转换超时"
    except Exception as e:
        return "[ERROR] " + str(e)
    return "[ERROR] LibreOffice 转换未生成文件"

def read_doc_binary(filepath):
    try:
        with open(filepath, 'rb') as f:
            data = f.read()
        if data[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':
            result = []
            for byte in data:
                if 32 <= byte <= 126 or byte in (9, 10, 13, 133):
                    result.append(chr(byte))
                else:
                    result.append(' ')
            text = ''.join(result)
            text = re.sub(r' {3,}', ' ', text)
            text = re.sub(r'\n{3,}', '\n\n', text)
            text = re.sub(r'[^\u4e00-\u9fff\u0000-\u007f\s\d\.,;:：；（）()【】《》""\'\'、，。；：！？\-–——\uff00-\uffef]', ' ', text)
            return text.strip()[:8000]
        return "[ERROR] 不是有效的 Word 文档"
    except Exception as e:
        return "[ERROR] " + str(e)

def detect_clauses(text):
    rules = [
        {
            'type': '廉洁/反腐条款',
            'keywords': ['廉洁', '反腐', '商业行为准则', '反腐败', '廉洁从业', '反腐承诺', '廉洁协议'],
            'weight': ['廉洁', '反腐', '反腐败', '商业行为准则'],
        },
        {
            'type': '违约金条款',
            'keywords': ['违约金', '违约赔偿', '违约责任', '违约方应向守约方支付', '每日按', '逾期付款违约金', '违约金的计算'],
            'weight': ['违约金', '违约赔偿', '违约责任'],
        },
        {
            'type': '付款条款',
            'keywords': ['付款方式', '付款时间', '付款期限', '结算方式', '结算周期', '支付方式', '账期', '付款条件', '首付', '尾款', '预付款', '全款', '分批付款'],
            'weight': ['付款方式', '付款时间', '付款期限', '结算方式', '账期'],
        },
        {
            'type': '保密条款',
            'keywords': ['保密义务', '保密责任', '商业秘密', '保密信息', '不得泄露', '不得向第三方披露', '保密协议', '保密条款', '信息保护', '知识产权'],
            'weight': ['保密义务', '保密责任', '商业秘密', '保密信息'],
        },
    ]

    results = []
    for rule in rules:
        matched = []
        all_kw = rule['weight'] + [k for k in rule['keywords'] if k not in rule['weight']]
        for kw in all_kw:
            try:
                pattern = re.compile(re.escape(kw), re.IGNORECASE)
            except:
                continue
            for m in pattern.finditer(text):
                pos = m.start()
                start = max(0, pos - 100)
                end = min(len(text), pos + len(kw) + 100)
                snippet = text[start:end].replace('\n', ' ')
                if len(snippet) > 10:
                    matched.append({'text': snippet, 'keyword': kw, 'pos': pos})
            if len(matched) >= 5:
                break
        seen = set()
        unique = []
        for item in matched:
            bucket = item['pos'] // 200
            if bucket not in seen:
                seen.add(bucket)
                unique.append(item)
        results.append({'type': rule['type'], 'items': unique[:5]})
    return results

def perform_audit(text, filename):
    results = []
    issues = []

    if text.startswith("[ERROR") or text.startswith("[WARNING"):
        results.append({'item': '文件读取', 'status': 'ERROR', 'detail': text[:200], 'suggestion': '请另存为 .docx 格式后重新审核'})
        return results, []

    if not text or not text.strip():
        results.append({'item': '文件读取', 'status': 'ERROR', 'detail': '提取内容为空', 'suggestion': '可能是扫描件或加密文件'})
        return results, []

    results.append({'item': '文件读取', 'status': 'PASS', 'detail': '成功提取 %d 字符' % len(text)})

    # 合同编号
    ht = re.search(r'HT(\d+)', text, re.I)
    if ht:
        results.append({'item': '合同编号', 'status': 'PASS', 'detail': 'HT' + ht.group(1)})
    else:
        issues.append({'item': '合同编号', 'status': 'WARN', 'detail': '未找到合同编号（格式：HT+数字）', 'suggestion': '补充合同编号'})

    # 甲方名称
    jf = re.search(r'甲\s*方[：:]\s*(.{2,50}?公司.{0,20}?)[\n\s]{2,}?(?:地址|电话|邮编|合同编号|开户|$)', text)
    if jf:
        name = re.sub(r'（[^）]*）', '', jf.group(1)).strip()
    else:
        # 备选：甲方：后面跟多行内容，取第一行非空
        jf2 = re.search(r'甲\s*方[：:]\s*([^\n]{2,60})', text)
        name = jf2.group(1).strip() if jf2 else None
        if name:
            name = re.sub(r'（[^）]*）', '', name).strip()
    if name and len(name) >= 2:
        results.append({'item': '甲方名称', 'status': 'PASS', 'detail': name})
    else:
        issues.append({'item': '甲方名称', 'status': 'WARN', 'detail': '无法识别甲方名称', 'suggestion': '检查合同格式'})

    # 乙方名称
    yf = re.search(r'乙\s*方[：:]\s*([^\n\s，,]{2,60})', text)
    if yf:
        name = re.sub(r'（[^）]*）', '', yf.group(1)).strip()
        std_name = STANDARDS.get('乙方名称', '')
        if std_name and (std_name in name or name in std_name):
            results.append({'item': '乙方名称', 'status': 'PASS', 'detail': name})
        elif name:
            results.append({'item': '乙方名称', 'status': 'INFO', 'detail': '检测到：' + name})
        else:
            issues.append({'item': '乙方名称', 'status': 'WARN', 'detail': '未找到乙方名称', 'suggestion': '检查合同格式'})
    else:
        issues.append({'item': '乙方名称', 'status': 'WARN', 'detail': '未找到乙方名称', 'suggestion': '检查合同格式'})

    # 乙方地址
    if STANDARDS['乙方地址'] in text:
        results.append({'item': '乙方地址', 'status': 'PASS', 'detail': STANDARDS['乙方地址']})
    else:
        issues.append({'item': '乙方地址', 'status': 'WARN', 'detail': '地址未匹配标准值', 'suggestion': '确认为' + STANDARDS['乙方地址']})

    # 乙方电话
    if STANDARDS.get('乙方电话') and STANDARDS['乙方电话'] in text:
        results.append({'item': '乙方电话', 'status': 'PASS', 'detail': STANDARDS['乙方电话']})
    else:
        issues.append({'item': '乙方电话', 'status': 'WARN', 'detail': '电话未匹配标准值', 'suggestion': '确认为' + STANDARDS['乙方电话']})

    # 收款银行
    bank_ok = True
    bd = []
    if STANDARDS['收款户名'] in text:
        bd.append('户名: ✓')
    else:
        bank_ok = False
        bd.append('户名: 未找到')
    if STANDARDS['收款开户行'] in text:
        bd.append('开户行: ✓')
    else:
        matched_b = next((v for v in BANK_VARIANTS if v in text), None)
        if matched_b:
            bd.append('开户行: ' + matched_b + ' ✓')
        else:
            bank_ok = False
            bd.append('开户行: 未匹配')

    found_acct = None
    # 收款账号：搜索"账号："或"账号"后的数字序列
    acct_std = STANDARDS.get('收款账号', '')
    for pat in [
        r'账\s*号[：:]\s*([\d\s]{15,25})',
        r'账号[：:]\s*([\d\s]{15,25})',
    ]:
        m = re.search(pat, text)
        if m:
            found_acct = m.group(1) if len(m.groups()) > 0 else m.group(0)
            break
    if found_acct:
        clean = re.sub(r'\s', '', found_acct)
        if clean == STANDARDS['收款账号']:
            bd.append('账号: ✓')
        else:
            bank_ok = False
            bd.append('账号: ⚠ 未匹配（' + found_acct.strip() + ' -> ' + clean + '）')
    else:
        bank_ok = False
        bd.append('账号: 未找到')

    if bank_ok:
        results.append({'item': '收款银行信息', 'status': 'PASS', 'detail': '；'.join(bd)})
    else:
        results.append({'item': '收款银行信息', 'status': 'WARN', 'detail': '；'.join(bd)})
        issues.append({'item': '收款银行信息', 'status': 'WARN', 'detail': '收款信息未完全匹配标准', 'suggestion': '核对收款账户'})

    # 签署日期
    tail = text[-1500:]
    def find_date(prefix):
        for p in [
            prefix + r'[^\n]*?\n\s*日\s*期[：:]\s*(\S+)',
            prefix + r'[^\n]*?[\n\s]*日\s*期[：:]\s*(\d{4}[-]\d{1,2}[-]\d{1,2}|\d{4}年\d{1,2}月\d{1,2}日)',
        ]:
            m = re.search(p, tail)
            if m:
                return m.group(1)
        m = re.search(prefix + r'.{0,100}?(\d{4}-\d{2}-\d{2})', tail)
        return m.group(1) if m else None

    jia_date = find_date('甲\s*方')
    yi_date = find_date('乙\s*方')

    def is_valid(d): return d and (bool(re.match(r'\d{4}-\d{2}-\d{2}', d)) or bool(re.match(r'\d{4}年\d{2}月\d{2}日', d)))
    def is_empty(d): return not d or len(str(d).strip()) < 5

    date_ok = True
    dd = []
    if is_empty(jia_date): date_ok = False; dd.append('甲方: 未填写')
    elif is_valid(jia_date): dd.append('甲方: ' + jia_date + ' ✓')
    else: date_ok = False; dd.append('甲方: ' + str(jia_date) + ' ⚠')
    if is_empty(yi_date): date_ok = False; dd.append('乙方: 未填写')
    elif is_valid(yi_date): dd.append('乙方: ' + yi_date + ' ✓')
    else: date_ok = False; dd.append('乙方: ' + str(yi_date) + ' ⚠')

    if date_ok:
        results.append({'item': '签署日期', 'status': 'PASS', 'detail': '；'.join(dd)})
    else:
        results.append({'item': '签署日期', 'status': 'WARN', 'detail': '；'.join(dd)})
        issues.append({'item': '签署日期', 'status': 'WARN', 'detail': '存在未填写或格式异常的日期', 'suggestion': '补充甲乙双方签署日期'})

    # 服务金额
    amt = re.search(r'[¥￥]?\s*([\d,]+\.?\d*)\s*元', text)
    if amt:
        results.append({'item': '服务金额', 'status': 'INFO', 'detail': '检测到金额: ¥' + amt.group(1).replace(',', '')})

    # 合并
    seen = {r['item'] for r in results}
    for iss in issues:
        if iss['item'] not in seen:
            results.append(iss)
            seen.add(iss['item'])

    clauses = detect_clauses(text)
    return results, clauses

@app.route('/audit', methods=['POST'])
def audit():
    if 'file' not in request.files:
        return jsonify({'error': '没有接收到文件'}), 400
    f = request.files['file']
    if not f.filename:
        return jsonify({'error': '文件名为空'}), 400
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in ('.docx', '.doc', '.pdf'):
        return jsonify({'error': '仅支持 .docx / .doc / .pdf 格式'}), 400

    fd, tmppath = tempfile.mkstemp(suffix=ext)
    os.close(fd)
    try:
        f.save(tmppath)
        text = extract_text(tmppath)
        results, clauses = perform_audit(text, f.filename)
        errors = sum(1 for r in results if r['status'] == 'ERROR')
        warns = sum(1 for r in results if r['status'] == 'WARN')
        overall = 'PASS' if errors == 0 and warns == 0 else ('ERROR' if errors > 0 else 'WARN')
        return jsonify({
            'filename': f.filename,
            'overall': overall,
            'results': results,
            'clauses': clauses,
            'text': text,  # 返回全文供前端 AI 使用
        })
    finally:
        try: os.remove(tmppath)
        except: pass

@app.route('/extract', methods=['POST'])
def extract_only():
    """仅提取文本，不做审核（给前端 AI 使用）"""
    if 'file' not in request.files:
        return jsonify({'error': '没有接收到文件'}), 400
    f = request.files['file']
    if not f.filename:
        return jsonify({'error': '文件名为空'}), 400
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in ('.docx', '.doc', '.pdf'):
        return jsonify({'error': '仅支持 .docx / .doc / .pdf 格式'}), 400

    fd, tmppath = tempfile.mkstemp(suffix=ext)
    os.close(fd)
    try:
        f.save(tmppath)
        text = extract_text(tmppath)
        return jsonify({
            'filename': f.filename,
            'text': text,
            'size': len(text)
        })
    finally:
        try: os.remove(tmppath)
        except: pass


@app.route('/convert', methods=['POST'])
def convert_to_docx():
    """Convert .doc to .docx, preserving original formatting.
    Tries Word COM first, then LibreOffice."""
    if 'file' not in request.files:
        return jsonify({'error': '没有接收到文件'}), 400
    f = request.files['file']
    if not f.filename:
        return jsonify({'error': '文件名为空'}), 400
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in ('.doc',):
        return jsonify({'error': '仅 .doc 文件可转换'}), 400

    fd, docpath = tempfile.mkstemp(suffix='.doc')
    os.close(fd)
    docxpath = docpath.replace('.doc', '.docx')

    try:
        f.save(docpath)

        # Method 1: Word COM via PowerShell (best quality)
        abs_doc = os.path.abspath(docpath).replace('\\', '\\\\')
        abs_docx = os.path.abspath(docxpath).replace('\\', '\\\\')
        ps = f'''
$ErrorActionPreference = 'Stop'
try {{
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $word.DisplayAlerts = 0
    $doc = $word.Documents.Open('{abs_doc}')
    $doc.SaveAs([ref]'{abs_docx}', [ref]16)
    $doc.Close()
    $word.Quit()
    Write-Output 'OK'
}} catch {{
    Write-Output "ERROR: $_"
}}
'''
        proc = subprocess.run(
            ['powershell', '-NoProfile', '-ExecutionPolicy', 'Bypass', '-Command', ps],
            capture_output=True, timeout=60, encoding='utf-8', errors='replace'
        )
        if proc.stdout.strip() == 'OK' and os.path.exists(docxpath) and os.path.getsize(docxpath) > 1000:
            with open(docxpath, 'rb') as df:
                docx_data = df.read()
            return docx_data, 200, {'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                      'Content-Disposition': f'attachment; filename="{os.path.splitext(f.filename)[0]}.docx"'}

        # Method 2: LibreOffice
        soffice_paths = [
            r'C:\Program Files\LibreOffice\program\soffice.exe',
            r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
            'soffice',
        ]
        soffice = None
        for p in soffice_paths:
            if os.path.exists(p):
                soffice = p
                break
            if p == 'soffice':
                try:
                    r = subprocess.run(['soffice', '--version'], capture_output=True, timeout=5)
                    if r.returncode == 0:
                        soffice = 'soffice'
                        break
                except:
                    pass

        if soffice:
            tmpdir = tempfile.gettempdir()
            subprocess.run([soffice, '--headless', '--convert-to', 'docx', '--outdir', tmpdir, docpath],
                           capture_output=True, timeout=30)
            base = os.path.splitext(os.path.basename(docpath))[0]
            result_path = os.path.join(tmpdir, base + '.docx')
            if os.path.exists(result_path) and os.path.getsize(result_path) > 1000:
                with open(result_path, 'rb') as df:
                    docx_data = df.read()
                os.remove(result_path)
                return docx_data, 200, {'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                          'Content-Disposition': f'attachment; filename="{os.path.splitext(f.filename)[0]}.docx"'}

        return jsonify({'error': '无法转换 .doc 文件。请安装 Microsoft Word 或 LibreOffice，或手动将 .doc 另存为 .docx 后重新上传。'}), 500

    except subprocess.TimeoutExpired:
        return jsonify({'error': '转换超时'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        try: os.remove(docpath)
        except: pass
        if os.path.exists(docxpath):
            try: os.remove(docxpath)
            except: pass


@app.route('/rules', methods=['GET', 'POST'])
def manage_rules():
    """规则同步端点：GET 获取服务端规则，POST 保存规则到服务端"""
    if request.method == 'GET':
        return jsonify({
            'rules': CFG.get('audit_rules', []),
            'standards': CFG.get('基本信息', DEFAULT_CONFIG['基本信息']),
            'bank_variants': CFG.get('开户行变体', DEFAULT_CONFIG['开户行变体']),
            'clause_checks': CFG.get('条款检查', DEFAULT_CONFIG['条款检查'])
        })

    if request.method == 'POST':
        data = request.get_json(silent=True) or {}
        if 'rules' in data:
            CFG['audit_rules'] = data['rules']
        if 'standards' in data:
            CFG['基本信息'] = data['standards']
        if 'bank_variants' in data:
            CFG['开户行变体'] = data['bank_variants']
        if 'clause_checks' in data:
            CFG['条款检查'] = data['clause_checks']
        save_config(CFG)
        # 重新加载全局变量
        global STANDARDS, BANK_VARIANTS
        STANDARDS = CFG.get('基本信息', DEFAULT_CONFIG['基本信息'])
        BANK_VARIANTS = CFG.get('开户行变体', DEFAULT_CONFIG['开户行变体'])
        return jsonify({'status': 'ok', 'message': '规则已保存'})


@app.route('/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok', 'version': '3.0'})


if __name__ == '__main__':
    print('=' * 50)
    print('  合同审核工具 v3 — 后端服务')
    print('  端口: 5577')
    print('  端点: POST /audit | POST /extract | GET/POST /rules | GET /health')
    print('=' * 50)
    app.run(host='127.0.0.1', port=5577, debug=False)
